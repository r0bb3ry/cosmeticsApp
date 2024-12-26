import sys
import matplotlib

matplotlib.use('Qt5Agg')
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QPushButton, QComboBox, QLabel, QTabWidget
from PyQt5 import uic
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import matplotlib.pyplot as plt
import mysql.connector
from datetime import datetime, date
import matplotlib.dates as mdates
from PyQt5.QtCore import Qt
from docx import Document
from docx.shared import Inches
import os


class StatsWindow(QWidget):
    def __init__(self, user_data=None):
        super().__init__()
        uic.loadUi('GrafWind.ui', self)
        self.user_data = user_data  # сохраняем user_data
        self.mydb = None
        self.connect_to_db()

        # Виджеты с первой вкладки (графики)
        self.frame = self.findChild(QWidget, "frame")
        # self.pushButton = self.findChild(QPushButton, "pushButton") убираем кнопку
        self.comboBox = self.findChild(QComboBox, "comboBox")
        self.comboBox_2 = self.findChild(QComboBox, "comboBox_2")
        self.comboBox_3 = self.findChild(QComboBox, "comboBox_3")

        # Виджеты со второй вкладки (статистика)
        self.comboBox_5 = self.findChild(QComboBox, "comboBox_5")
        self.comboBox_4 = self.findChild(QComboBox, "comboBox_4")
        self.frame_2 = self.findChild(QWidget, "frame_2")
        self.comboBox_6 = self.findChild(QComboBox, "comboBox_6")
        # self.comboBox_month = self.findChild(QComboBox, "comb_onth") убираем выбор месяца со второго таба

        # Вкладки
        self.tabs = self.findChild(QTabWidget, "tabWidget")

        # Инициализация графика
        self.figure = Figure(figsize=(8, 6), dpi=100)
        self.axes = self.figure.add_subplot(111)
        self.canvas = FigureCanvas(self.figure)

        self.figure_2 = Figure(figsize=(8, 6), dpi=100)
        self.axes_2 = self.figure_2.add_subplot(111)
        self.canvas_2 = FigureCanvas(self.figure_2)

        # Добавление графика во frame
        layout = QVBoxLayout(self.frame)
        layout.addWidget(self.canvas)

        layout_2 = QVBoxLayout(self.frame_2)
        layout_2.addWidget(self.canvas_2)

        # self.pushButton.clicked.connect(self.update_plot) убираем сигнал по кнопке

        # Загрузка списков
        self.load_product_types()
        self.load_months()
        self.load_time_periods()
        self.load_employees()

        self.comboBox_2.currentIndexChanged.connect(self.toggle_month_combobox)
        self.comboBox_4.currentIndexChanged.connect(self.update_plot_2)
        self.comboBox_5.currentIndexChanged.connect(self.update_plot_2)
        self.comboBox_6.currentIndexChanged.connect(self.update_plot_2)
        # self.comboBox_month.currentIndexChanged.connect(self.update_plot_2) убираем сигнал выбора месяца
        self.comboBox.currentIndexChanged.connect(self.update_plot)
        self.comboBox_3.currentIndexChanged.connect(
            self.update_plot)  # Добавлено - сигнал изменения combobox для обновления графика

        self.pushButton_report_tab1 = self.findChild(QPushButton, "pushButton_report_tab1")
        self.pushButton_report_tab1.clicked.connect(self.create_report_tab1)

        self.pushButton_report_tab2 = self.findChild(QPushButton, "pushButton_report_tab2")
        self.pushButton_report_tab2.clicked.connect(self.create_report_tab2)

        self.update_plot()
        self.update_plot_2()

    def connect_to_db(self):
        try:
            self.mydb = mysql.connector.connect(
                host="localhost",
                user="root",
                password="2810300719",
                database="dianadb"
            )
            if self.mydb.is_connected():
                print("Connected to MySQL Database")
        except mysql.connector.Error as e:
            print(f"Error connecting to MySQL: {e}")
            sys.exit(1)

    def load_product_types(self):
        cursor = self.mydb.cursor()
        try:
            cursor.execute("SELECT product_type_name FROM product_types")
            product_types = [row[0] for row in cursor.fetchall()]
            self.comboBox.addItems(product_types)
            self.comboBox_6.addItems(["Все товары"] + product_types)
        except mysql.connector.Error as e:
            print(f"Error loading product types: {e}")
            sys.exit(1)
        finally:
            cursor.close()

    def load_months(self):
        months = ["Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
                  "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"]
        self.comboBox_3.addItems(months)
        # self.comboBox_month.addItems(months) удаляем заполнение месяцев

    def load_time_periods(self):
        time_periods = ["Месяц", "Год"]
        self.comboBox_2.addItems(time_periods)
        years = [str(year) for year in range(2020, datetime.now().year + 1)]
        self.comboBox_4.addItems(years)

    def load_employees(self):
        cursor = self.mydb.cursor()
        try:
            cursor.execute("SELECT first_name, last_name, employee_id FROM employees")
            employees = [f"{row[0]} {row[1]} ({row[2]})" for row in cursor.fetchall()]
            self.comboBox_5.addItems(["Общие"] + employees)
        except mysql.connector.Error as e:
            print(f"Error loading employees: {e}")
            sys.exit(1)
        finally:
            cursor.close()

    def toggle_month_combobox(self, index):
        selected_period = self.comboBox_2.itemText(index)
        if selected_period == "Месяц":
            self.comboBox_3.show()
        else:
            self.comboBox_3.hide()

    def update_plot(self):
        selected_type = self.comboBox.currentText()
        selected_period = self.comboBox_2.currentText()
        selected_month = self.comboBox_3.currentText() if selected_period == "Месяц" else ""

        self.axes.clear()  # Очищаем предыдущий график

        sales_data = self.get_sales_data(selected_type, selected_period, selected_month)

        if sales_data and 'dates' in sales_data and 'values' in sales_data:
            dates = [date.fromisoformat(str(d)) for d in sales_data['dates']]
            values = sales_data['values']

            self.axes.plot(dates, values, marker='o', linestyle='-', color='skyblue')
            self.axes.set_title(f'Продажи ({selected_type}, {selected_period}, {selected_month})')
            self.axes.set_xlabel('Дата')
            self.axes.set_ylabel('Количество')
            self.axes.grid(True, linestyle='--', alpha=0.6)
            self.axes.xaxis.set_major_locator(mdates.AutoDateLocator())
            self.axes.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
            plt.setp(self.axes.get_xticklabels(), rotation=45, ha='right')
            self.canvas.draw()
        else:
            self.axes.text(0.5, 0.5, 'Нет данных',
                           fontsize=12, horizontalalignment='center', verticalalignment='center')
            self.canvas.draw()

    def get_sales_data(self, product_type, period, selected_month):
        cursor = self.mydb.cursor()
        try:
            if period == "Месяц" and selected_month:
                month_num = self.comboBox_3.currentIndex() + 1
                sql = """
                      SELECT s.sale_date, SUM(s.quantity)
                    FROM sales s
                    JOIN products p ON s.product_id = p.product_id
                    JOIN product_types pt ON p.product_type_id = pt.product_type_id
                    WHERE YEAR(s.sale_date) = 2024
                    AND MONTH(s.sale_date) = %s
                    AND pt.product_type_name = %s
                    GROUP BY s.sale_date
                    ORDER BY s.sale_date
                 """
                cursor.execute(sql, (month_num, product_type))
            elif period == "Год":
                sql = """
                    SELECT s.sale_date, SUM(s.quantity)
                    FROM sales s
                    JOIN products p ON s.product_id = p.product_id
                    JOIN product_types pt ON p.product_type_id = pt.product_type_id
                    WHERE YEAR(s.sale_date) = 2024
                    AND pt.product_type_name = %s
                    GROUP BY s.sale_date
                    ORDER BY s.sale_date
                """
                cursor.execute(sql, (product_type,))
            results = cursor.fetchall()
            if not results:
                return None
            sale_dates = [row[0] for row in results]
            quantities = [row[1] for row in results]
            return {'dates': sale_dates, 'values': quantities}
        except mysql.connector.Error as e:
            print(f"Error getting sales data: {e}")
            return None
        finally:
            cursor.close()

    def update_plot_2(self):
        selected_type = self.comboBox_6.currentText()
        selected_year = self.comboBox_4.currentText()
        selected_employee = self.comboBox_5.currentText()
        # month_num = None # удаляем

        # if selected_period == "Месяц": #удаляем
        # if not self.comboBox_month.count():
        # print("Список месяцев не загружен!")
        # return
        # month_num = self.comboBox_month.currentIndex() + 1

        self.axes_2.clear()

        products_data = self.get_sales_data_2(selected_type, selected_year, selected_employee)

        if products_data:
            names = list(products_data.keys())
            values = list(products_data.values())
            self.axes_2.bar(names, values)
            self.axes_2.set_xlabel("Месяцы")
            self.axes_2.set_ylabel("Продажи")
            self.axes_2.set_title(f"Продажи ({selected_type}, {selected_year}, {selected_employee})")
            self.axes_2.set_xticks(names)
            self.canvas_2.draw()
        else:
            self.axes_2.text(0.5, 0.5, 'Нет данных',
                             fontsize=12, horizontalalignment='center', verticalalignment='center')
            self.canvas_2.draw()

    def get_sales_data_2(self, product_type, year, employee):
        cursor = self.mydb.cursor()
        try:
            if employee == "Общие":
                sql = """
                        SELECT MONTH(s.sale_date), SUM(s.quantity)
                        FROM sales s
                        JOIN products p ON s.product_id = p.product_id
                        JOIN product_types pt ON p.product_type_id = pt.product_type_id
                        WHERE YEAR(s.sale_date) = %s
                        AND (pt.product_type_name = %s OR %s = 'Все товары')
                        GROUP BY MONTH(s.sale_date)
                        ORDER BY MONTH(s.sale_date)
                    """
                cursor.execute(sql, (year, product_type, product_type))
            else:
                employee_id = employee.split("(")[1][:-1]
                sql = """
                            SELECT MONTH(s.sale_date), SUM(s.quantity)
                            FROM sales s
                            JOIN products p ON s.product_id = p.product_id
                            JOIN product_types pt ON p.product_type_id = pt.product_type_id
                            WHERE YEAR(s.sale_date) = %s
                            AND s.employee_id = %s
                            AND (pt.product_type_name = %s OR %s = 'Все товары')
                            GROUP BY MONTH(s.sale_date)
                            ORDER BY MONTH(s.sale_date)
                       """
                cursor.execute(sql, (year, employee_id, product_type, product_type))

            results = cursor.fetchall()
            if not results:
                return None

            months = ["Декабрь", "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь", "Июль", "Август", "Сентябрь",
                      "Октябрь", "Ноябрь"]
            data = {}
            for i, month_num in enumerate(range(12, 0, -1)):
                found = False
                for row in results:
                    if row[0] == month_num:
                        data[months[i]] = row[1]
                        found = True
                        break
                if not found:
                    data[months[i]] = 0
            return data

        except mysql.connector.Error as e:
            print(f"Error getting sales data: {e}")
            return None
        finally:
            cursor.close()

    def create_report_tab1(self):
        try:
            if self.user_data is None:
                print("Нет данных о пользователе")
                return

            doc = Document()
            doc.add_heading("Отчет о продажах (Первая вкладка)", level=1)
            first_name = self.user_data[1] if self.user_data[1] else "Неизвестно"
            last_name = self.user_data[2] if self.user_data[2] else "Неизвестно"
            selected_type = self.comboBox.currentText()
            selected_period = self.comboBox_2.currentText()
            selected_month = self.comboBox_3.currentText() if selected_period == "Месяц" else ""

            doc.add_paragraph(f"Отчет создан: {first_name} {last_name}")
            doc.add_paragraph(f"Тип товара: {selected_type}")
            doc.add_paragraph(f"Период: {selected_period}, {selected_month}")

            if self.axes.lines:
                doc.add_paragraph("График продаж за выбранный период")
                filename = "report_tab1_temp.png"
                self.canvas.figure.savefig(filename)
                doc.add_picture(filename, width=Inches(6))
                os.remove(filename)  # удаляем файл после вставки
            else:
                doc.add_paragraph("График продаж пуст")
            doc.save("report_tab1.docx")
            print("Отчет сохранен как report_tab1.docx")
        except Exception as e:
            print(f"Error creating report_tab1: {e}")

    def create_report_tab2(self):
        try:
            if self.user_data is None:
                print("Нет данных о пользователе")
                return

            doc = Document()
            doc.add_heading("Отчет о продажах (Вторая вкладка)", level=1)
            first_name = self.user_data[1] if self.user_data[1] else "Неизвестно"
            last_name = self.user_data[2] if self.user_data[2] else "Неизвестно"
            selected_type = self.comboBox_6.currentText()
            selected_year = self.comboBox_4.currentText()
            selected_employee = self.comboBox_5.currentText()

            doc.add_paragraph(f"Отчет создан: {first_name} {last_name}")
            doc.add_paragraph(f"Тип товара: {selected_type}")
            doc.add_paragraph(f"Год: {selected_year}, Сотрудник: {selected_employee}")

            if self.axes_2.patches:
                doc.add_paragraph("Столбчатая диаграмма продаж по месяцам")
                filename = "report_tab2_temp.png"
                self.canvas_2.figure.savefig(filename)
                doc.add_picture(filename, width=Inches(6))
                os.remove(filename)  # удаляем файл после вставки
            else:
                doc.add_paragraph("Столбчатая диаграмма продаж пуста")
            doc.save("report_tab2.docx")
            print("Отчет сохранен как report_tab2.docx")
        except Exception as e:
            print(f"Error creating report_tab2: {e}")

    def closeEvent(self, event):
        if self.mydb:
            self.mydb.close()
        event.accept()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = StatsWindow()
    window.show()
    sys.exit(app.exec_())